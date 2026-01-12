import re
import os
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def add_hyperlink(paragraph, text, url):
    # This gets access to the document.xml.rels file and gets a new relation id value
    part = paragraph.part
    r_id = part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True)

    # Create the w:hyperlink tag and add needed values
    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), r_id)

    # Create a w:r element and a new w:rPr element
    new_run = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')

    # Join all the xml elements together add add the required text to the w:r element
    new_run.append(rPr)
    new_run.text = text
    hyperlink.append(new_run)

    paragraph._p.append(hyperlink)
    
    return hyperlink

def md_to_docx(md_path, docx_path, results_dir):
    doc = Document()
    
    # Style configuration
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)

    with open(md_path, 'r', encoding='utf-8') as f:
        lines = f.readlines()

    in_table = False
    table_lines = []
    
    # Pre-compile regex
    img_pattern = re.compile(r'!\[(.*?)\]\((.*?)\)')
    header_pattern = re.compile(r'^(#{1,6})\s+(.*)')
    list_pattern = re.compile(r'^(\s*)[-*]\s+(.*)')

    i = 0
    while i < len(lines):
        line = lines[i].strip()
        
        # Skip empty lines unless we need spacing, but docx handles spacing well usually
        if not line:
            if in_table:
                # Process the table we just finished
                process_table(doc, table_lines)
                in_table = False
                table_lines = []
            i += 1
            continue

        # Detect Tables
        if line.startswith('|'):
            in_table = True
            table_lines.append(line)
            i += 1
            continue
        elif in_table:
            # End of table
            process_table(doc, table_lines)
            in_table = False
            table_lines = []
            # Don't increment i, process this line normally
            continue

        # Headers
        header_match = header_pattern.match(line)
        if header_match:
            level = len(header_match.group(1))
            text = header_match.group(2)
            doc.add_heading(text, level=min(level, 9))
            i += 1
            continue

        # Images
        img_match = img_pattern.match(line)
        if img_match:
            caption = img_match.group(1)
            path = img_match.group(2)
            
            # Fix path
            if path.startswith('/api/results/'):
                filename = path.replace('/api/results/', '')
                local_path = os.path.join(results_dir, filename)
            else:
                local_path = path

            if os.path.exists(local_path):
                try:
                    doc.add_picture(local_path, width=Inches(6))
                    last_paragraph = doc.paragraphs[-1] 
                    last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    
                    # Add caption
                    caption_para = doc.add_paragraph(caption)
                    caption_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    caption_run = caption_para.runs[0]
                    caption_run.italic = True
                    caption_run.font.size = Pt(9)
                except Exception as e:
                    doc.add_paragraph(f"[Error inserting image: {local_path} - {str(e)}]")
            else:
                doc.add_paragraph(f"[Image not found: {local_path}]")
            
            i += 1
            continue

        # Lists
        list_match = list_pattern.match(line)
        if list_match:
            indent = len(list_match.group(1))
            text = list_match.group(2)
            # Simplified list handling (bullet points)
            p = doc.add_paragraph(text, style='List Bullet')
            # Indentation handling could be added here if needed
            i += 1
            continue

        # Standard Paragraph
        # Handle simple bold/italic (very basic)
        p = doc.add_paragraph()
        process_inline_formatting(p, line)
        i += 1

    if in_table:
        process_table(doc, table_lines)

    doc.save(docx_path)
    print(f"Successfully saved to {docx_path}")

def process_table(doc, lines):
    # Basic pipe table parser
    # Remove alignment row (---)
    clean_lines = [l for l in lines if not set(l.strip().replace('|', '').replace('-', '').replace(':', '')) == set()]
    
    if not clean_lines:
        return

    # Determine columns
    header_row = [c.strip() for c in clean_lines[0].split('|') if c.strip()]
    cols = len(header_row)
    
    table = doc.add_table(rows=0, cols=cols)
    table.style = 'Table Grid'
    
    # Header
    row_cells = table.add_row().cells
    for j, text in enumerate(header_row):
        row_cells[j].text = text
        # Make bold
        for paragraph in row_cells[j].paragraphs:
            for run in paragraph.runs:
                run.font.bold = True

    # Data
    for line in clean_lines[1:]:
        data_row = [c.strip() for c in line.split('|') if c.strip() or c == '']
        # Handle mismatch cols (sometimes split creates empty start/end)
        # Filter purely empty strings if they resulted from leading/trailing pipes
        cells_data = [d for d in line.strip().split('|')]
        # Remove first and last if they are empty strings (common in | a | b | format)
        if cells_data[0] == '': cells_data.pop(0)
        if cells_data and cells_data[-1] == '': cells_data.pop(-1)
        
        if len(cells_data) != cols:
            # Fallback or truncate
            cells_data = cells_data[:cols]
        
        row_cells = table.add_row().cells
        for j, text in enumerate(cells_data):
            if j < len(row_cells):
                row_cells[j].text = text

def process_inline_formatting(paragraph, text):
    # Very basic parser for **bold** and *italic*
    # This is a bit fragile but works for simple cases
    # We will just write the text for now to avoid complexity in this snippet
    # verifying robust regex replacement for runs is complex
    
    # Simple bold ** processing
    parts = re.split(r'(\*\*.*?\*\*)', text)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        else:
            paragraph.add_run(part)

if __name__ == "__main__":
    md_file = r"c:\Users\LLOYD\Downloads\nasa data\assessment_deliverables\S2_Research_Paper_Basis.md"
    docx_file = r"c:\Users\LLOYD\Downloads\nasa data\assessment_deliverables\S2_Research_Paper.docx"
    results_dir = r"c:\Users\LLOYD\Downloads\nasa data\results"
    
    try:
        md_to_docx(md_file, docx_file, results_dir)
    except Exception as e:
        print(f"Error: {e}")
